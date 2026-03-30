"""Tabla dinámica de trabajadores sobre el DOCX ya cargado (después del reemplazo ZIP)."""

from __future__ import annotations

from copy import deepcopy
from typing import Any

from docx.document import Document as DocumentClass
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table


def _cell_text_lower(cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip().lower()


def _row_header_score(row) -> int:
    texts = [_cell_text_lower(c) for c in row.cells]
    joined = " ".join(texts)
    score = 0
    if "nombre" in joined and "trabaj" in joined:
        score += 2
    if "imss" in joined or "nss" in joined:
        score += 2
    if "actividad" in joined:
        score += 1
    if "tel" in joined or "emer" in joined:
        score += 1
    return score


def find_worker_table(doc: DocumentClass) -> Table | None:
    best: tuple[int, int, Table] | None = None
    for table in doc.tables:
        if len(table.rows) < 2 or len(table.columns) < 4:
            continue
        for row in table.rows:
            sc = _row_header_score(row)
            if sc >= 4:
                key = (sc, len(table.columns))
                if best is None or key > (best[0], best[1]):
                    best = (key[0], key[1], table)
    if best:
        return best[2]
    for table in doc.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 4:
            if _row_header_score(table.rows[0]) >= 3:
                return table
    for table in doc.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 4:
            return table
    return None


def worker_header_row_index(table: Table) -> int:
    for i, row in enumerate(table.rows):
        if _row_header_score(row) >= 4:
            return i
    return 0


def _delete_row(table: Table, idx: int) -> None:
    tr = table.rows[idx]._tr
    tr.getparent().remove(tr)


def _set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        p = cell.paragraphs[0]
        runs = list(p.runs)
        if runs:
            runs[0].text = text or ""
            for r in runs[1:]:
                r.text = ""
        else:
            p.add_run(text or "")
    else:
        cell.text = text or ""


def _mk_w_border(side: str, sz: str = "6") -> OxmlElement:
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "000000")
    return el


def _ensure_tbl_pr(tbl_el) -> OxmlElement:
    tbl_pr = tbl_el.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    return tbl_pr


def _tc_apply_black_borders_all_sides(tc_el) -> None:
    """Cada w:tc del XML (incl. celdas lógicas por gridSpan), no solo row.cells."""
    tc_pr = tc_el.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc_el.insert(0, tc_pr)
    old = tc_pr.find(qn("w:tcBorders"))
    if old is not None:
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        borders.append(_mk_w_border(side, "6"))
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is not None:
        tc_w.addnext(borders)
    else:
        tc_pr.insert(0, borders)


def _force_worker_table_borders_for_pdf(table: Table) -> None:
    tbl_el = table._tbl
    tbl_pr = _ensure_tbl_pr(tbl_el)

    ts = tbl_pr.find(qn("w:tblStyle"))
    if ts is not None:
        tbl_pr.remove(ts)
    look = tbl_pr.find(qn("w:tblLook"))
    if look is not None:
        tbl_pr.remove(look)

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "autofit")

    old_tb = tbl_pr.find(qn("w:tblBorders"))
    if old_tb is not None:
        tbl_pr.remove(old_tb)
    tb = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tb.append(_mk_w_border(side, "6"))
    tbl_pr.append(tb)

    for tr in tbl_el:
        if tr.tag != qn("w:tr"):
            continue
        for tc in tr:
            if tc.tag != qn("w:tc"):
                continue
            _tc_apply_black_borders_all_sides(tc)


def fill_worker_table(doc: DocumentClass, workers: list[dict[str, Any]]) -> None:
    table = find_worker_table(doc)
    if table is None:
        return

    hi = worker_header_row_index(table)
    if hi + 1 >= len(table.rows):
        return

    prototype_tr = deepcopy(table.rows[hi + 1]._tr)

    while len(table.rows) > hi + 1:
        _delete_row(table, len(table.rows) - 1)

    rows_data = workers if workers else [{"nombre": "—", "imss": "", "actividad": "", "tel": ""}]

    for w in rows_data:
        new_tr = deepcopy(prototype_tr)
        table._tbl.append(new_tr)
        row = table.rows[-1]
        vals = [
            str(w.get("nombre") or ""),
            str(w.get("imss") or ""),
            str(w.get("actividad") or ""),
            str(w.get("tel") or ""),
        ]
        for i, val in enumerate(vals):
            if i < len(row.cells):
                _set_cell_text(row.cells[i], val or "")

    _force_worker_table_borders_for_pdf(table)
