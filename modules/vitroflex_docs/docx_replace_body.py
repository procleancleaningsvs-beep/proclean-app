"""Sustitución de placeholders en cuerpo, tablas, encabezados y pies (texto fusionado por python-docx)."""

from __future__ import annotations

from docx.document import Document as DocumentClass
from docx.text.paragraph import Paragraph


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Reemplaza el texto del párrafo conservando el primer run; limpia el resto."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def replace_in_paragraph(paragraph: Paragraph, mapping: dict[str, str]) -> None:
    full = paragraph.text
    if not full.strip():
        return
    new = full
    for k, v in mapping.items():
        if k in new:
            new = new.replace(k, v)
    if new != full:
        _set_paragraph_text(paragraph, new)


def replace_in_document(doc: DocumentClass, mapping: dict[str, str]) -> None:
    """Aplica mapping en párrafos, tablas, encabezados y pies."""
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)

    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p, mapping)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p, mapping)

        for p in section.footer.paragraphs:
            replace_in_paragraph(p, mapping)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p, mapping)
